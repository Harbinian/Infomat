import os, sys
from docx import Document

sys.stdout.reconfigure(encoding='utf-8')

base = r"E:\CA001\Infomat\docs\norms\运维安环部业务资料"

# The 14 target programs + 2 reference files
target_dirs = {
    # Process 1: 职业病危害识别申报与告知
    "职业病危害项目申报管理程序": "P1_申报",
    "职业病危害警示与告知管理程序": "P1_警示告知",
    # Process 2: 职业病防护与劳动防护
    "职业病防护设施管理程序": "P2_防护设施",
    "劳动防护用品管理程序": "P2_劳防用品",
    # Process 3: 职业危害监测与健康监护
    "职业病危害监测及评价管理程序": "P3_监测评价",
    "职业健康监护管理程序": "P3_健康监护",
    "职业健康档案管理程序": "P3_健康档案",
    "职业性噪声聋预警管理程序": "P3_噪声聋",
    # Process 4: 职业健康安全体系运行
    "职业健康安全目标指标管理程序": "P4_目标指标",
    "职业健康安全培训教育控制程序": "P4_培训教育",
    "职业健康安全风险管控管理程序": "P4_风险管控",
    "职业健康安全应急准备与响应控制程序": "P4_应急准备",
    "职业健康安全监视与测量控制管理程序": "P4_监视测量",
    "职业健康安全改进管理程序": "P4_改进",
    # Reference
    "职业卫生管理规则": "REF_规则",
    "职业病危害事故应急管理程序": "REF_事故应急",
}

for dirname, tag in target_dirs.items():
    dirpath = os.path.join(base, dirname)
    if not os.path.isdir(dirpath):
        print(f"\n{'='*80}\nNOT FOUND [{tag}]: {dirname}\n{'='*80}")
        continue

    files = os.listdir(dirpath)
    docx_files = [f for f in files if f.endswith('.docx') and not f.startswith('~$')]

    if not docx_files:
        print(f"\n{'='*80}\nNO DOCX [{tag}]: {dirname} — files: {files}\n{'='*80}")
        continue

    fpath = os.path.join(dirpath, docx_files[0])
    doc = Document(fpath)

    print(f"\n{'='*80}")
    print(f"[{tag}] DIR: {dirname}")
    print(f"FILE: {docx_files[0]}")
    print(f"PARAGRAPHS: {len(doc.paragraphs)}, TABLES: {len(doc.tables)}")
    print(f"{'='*80}")

    # Print all substantive paragraphs
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            style = p.style.name if p.style else ""
            print(f"  P[{i}]({style}): {text}")

    # Print all tables in full
    for ti, table in enumerate(doc.tables):
        print(f"\n  --- TABLE {ti} ({len(table.rows)} rows x {len(table.columns)} cols) ---")
        for ri, row in enumerate(table.rows):
            cells = [cell.text.strip().replace('\n', ' | ') for cell in row.cells]
            print(f"    ROW[{ri}]: {' || '.join(cells)}")
