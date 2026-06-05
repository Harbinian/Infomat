#!/usr/bin/env python3
"""Extract text from all 节能双碳 & 能源管理 DOCX files with structure info."""
import os, sys
from docx import Document

BASE = r"E:\CA001\Infomat\docs\norms\运维安环部业务资料"

FILES = [
    ("能源评审与年度节能双碳工作策划程序", "能源评审与年度节能双碳工作策划程序.docx"),
    ("动力系统节能运行管理程序", "动力系统节能运行管理程序.docx"),
    ("重点用能工序节能运行管理程序", "重点用能工序节能运行管理程序.docx"),
    ("能源消耗与能源监督管理程序", "能源消耗与能源监督管理程序.docx"),
    ("节能双碳监督检查管理程序", "节能双碳监督检查管理程序.docx"),
    ("能源碳排放统计分析与绩效管理程序", "能源碳排放统计分析与绩效管理程序.docx"),
    ("能源计量管理程序", "能源计量管理程序.docx"),
    ("淘汰落后用能管理程序", "淘汰落后用能管理程序.docx"),
    ("能源内部审核管理程序", "能源内部审核管理程序.docx"),
    ("能源管理评审管理程序", "能源管理评审管理程序.docx"),
    ("环境保护\"三同时\"和节能审查管理程序", "环境保护\"三同时\"和节能审查管理程序.docx"),
]

OUT = os.path.join(BASE, "_energy_texts.txt")

def para_info(p, idx):
    """Return (style_name, is_bold, text, list_info) for a paragraph."""
    text = p.text.strip()
    style = p.style.name if p.style else "None"
    # Detect bold runs
    bold_runs = [r.bold for r in p.runs if r.bold]
    is_bold = any(bold_runs) if bold_runs else False
    # Try to detect list numbering
    numPr = p._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
    has_list = "LIST" if numPr is not None else ""
    return (style, is_bold, text, has_list)

def extract_table(table):
    """Extract table as formatted text."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace('\n', ' | ') for cell in row.cells]
        rows.append(" | ".join(cells))
    return "\n".join(rows)

with open(OUT, 'w', encoding='utf-8') as out:
    for dir_name, file_name in FILES:
        path = os.path.join(BASE, dir_name, file_name)
        if not os.path.exists(path):
            out.write(f"\n{'='*80}\nFILE NOT FOUND: {path}\n{'='*80}\n\n")
            continue

        out.write(f"\n{'='*80}\n")
        out.write(f"FILE: {dir_name}/{file_name}\n")
        out.write(f"{'='*80}\n\n")

        doc = Document(path)

        out.write("--- PARAGRAPHS (with style) ---\n\n")
        for i, p in enumerate(doc.paragraphs):
            style, is_bold, text, has_list = para_info(p, i)
            if not text and not has_list:
                continue
            flags = []
            if is_bold:
                flags.append("BOLD")
            if has_list:
                flags.append("LIST")
            flag_str = f" [{', '.join(flags)}]" if flags else ""
            out.write(f"[{i}] <{style}>{flag_str} {text}\n")

        out.write(f"\n--- TABLES ({len(doc.tables)} total) ---\n\n")
        for ti, table in enumerate(doc.tables):
            out.write(f"\n>>> TABLE {ti}:\n")
            out.write(extract_table(table))
            out.write("\n")

print(f"Extraction complete. Output: {OUT}")
