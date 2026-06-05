# -*- coding: utf-8 -*-
"""Extract text from all equipment DOCX files for business behavior analysis."""
import os
from docx import Document

BASE = r"E:\CA001\Infomat\docs\norms\运维安环部业务资料\设备管理制度"
OUTPUT = r"E:\CA001\Infomat\docs\norms\运维安环部业务资料\_equipment_extracted.txt"

FILES = [
    ("1_计划策划", "设备设施计划策划管理程序.docx"),
    ("2a_使用", "设备设施使用管理程序.docx"),
    ("2b_操作证", "设备操作证管理标准.docx"),
    ("2c_巡检", "巡检实施管理标准.docx"),
    ("2d_激光跟踪仪", "激光跟踪仪的测量管理标准.docx"),
    ("3a_维护保养", "设备设施维护保养管理程序.docx"),
    ("3b_维修", "设备设施维修管理程序.docx"),
    ("3c_保养优化", "设备设施保养优化管理程序.docx"),
    ("3d_TPM", "设备设施TPM应用管理标准.docx"),
    ("4_备件", "设备设施备件管理程序.docx"),
    ("5a_事故处理", "设备设施事故处理管理程序.docx"),
    ("5b_搬迁", "设备设施搬迁管理程序.docx"),
]

def extract_paragraphs(doc):
    lines = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            style = para.style.name if para.style else "Normal"
            lines.append(f"  P[{i}] [{style}] {text}")
    return lines

def extract_tables(doc):
    lines = []
    for ti, table in enumerate(doc.tables):
        lines.append(f"\n  --- TABLE {ti+1} ({len(table.rows)} rows x {len(table.columns)} cols) ---")
        for ri, row in enumerate(table.rows):
            cells = [cell.text.strip().replace('\n', ' | ') for cell in row.cells]
            lines.append(f"  T[{ti}].R[{ri}]: {' || '.join(cells)}")
    return lines

with open(OUTPUT, 'w', encoding='utf-8') as out:
    for label, fname in FILES:
        fpath = os.path.join(BASE, fname)
        out.write(f"\n{'='*80}\n")
        out.write(f"FILE: {label} -- {fname}\n")
        out.write(f"{'='*80}\n")

        if not os.path.exists(fpath):
            out.write("  [FILE NOT FOUND]\n")
            continue

        fsize = os.path.getsize(fpath)
        if fsize == 0:
            out.write("  [FILE IS EMPTY - 0 bytes]\n")
            continue

        try:
            doc = Document(fpath)
            out.write(f"  File size: {fsize} bytes\n")
            out.write(f"  Paragraphs: {len(doc.paragraphs)}\n")
            out.write(f"  Tables: {len(doc.tables)}\n\n")

            out.write("  === PARAGRAPHS ===\n")
            for line in extract_paragraphs(doc):
                out.write(line + "\n")

            out.write("\n  === TABLES ===\n")
            for line in extract_tables(doc):
                out.write(line + "\n")

        except Exception as e:
            out.write(f"  [ERROR reading file: {e}]\n")

print(f"Done. Output written to: {OUTPUT}")
