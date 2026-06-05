import os
from docx import Document

base = r"E:\CA001\Infomat\docs\norms\运维安环部业务资料"

files = {
    # Process 1
    "职业病危害项目申报管理程序.docx": "P1",
    "职业病危害警示与告知管理程序.docx": "P1",
    # Process 2
    "职业病防护设施管理程序.docx": "P2",
    "劳动防护用品管理程序.docx": "P2",
    # Process 3
    "职业病危害监测及评价管理程序.docx": "P3",
    "职业健康监护管理程序.docx": "P3",
    "职业健康档案管理程序.docx": "P3",
    "职业性噪声聋预警管理程序.docx": "P3",
    # Process 4
    "职业健康安全目标指标管理程序.docx": "P4",
    "职业健康安全培训教育控制程序.docx": "P4",
    "职业健康安全风险管控管理程序.docx": "P4",
    "职业健康安全应急准备与响应控制程序.docx": "P4",
    "职业健康安全监视与测量控制管理程序.docx": "P4",
    "职业健康安全改进管理程序.docx": "P4",
    # Extra reference
    "职业卫生管理规则.docx": "REF",
    "职业病危害事故应急管理程序.docx": "REF",
}

for fname, group in files.items():
    fpath = os.path.join(base, fname)
    if not os.path.exists(fpath):
        print(f"\n{'='*80}\nMISSING [{group}]: {fname}\n{'='*80}")
        continue
    doc = Document(fpath)

    # Extract file number from headers/footers or first page
    header_text = ""
    for section in doc.sections:
        header = section.header
        if header:
            for p in header.paragraphs:
                header_text += p.text + "\n"

    print(f"\n{'='*80}")
    print(f"[{group}] FILE: {fname}")
    print(f"PARAGRAPHS: {len(doc.paragraphs)}")
    print(f"TABLES: {len(doc.tables)}")
    print(f"{'='*80}")

    # Print header info if available
    if header_text.strip():
        print(f"--- HEADER ---\n{header_text.strip()}")

    # Print first 10 paragraphs for document number / title info
    for i, p in enumerate(doc.paragraphs[:10]):
        if p.text.strip():
            print(f"  P[{i}]: {p.text.strip()}")

    # Print all paragraphs with content
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text and len(text) > 5:
            style = p.style.name if p.style else ""
            print(f"  P[{i}]({style}): {text}")

    # Print tables
    for ti, table in enumerate(doc.tables):
        print(f"\n  --- TABLE {ti} ({len(table.rows)} rows x {len(table.columns)} cols) ---")
        for ri, row in enumerate(table.rows):
            cells = [cell.text.strip().replace('\n', ' | ') for cell in row.cells]
            print(f"    ROW[{ri}]: {' || '.join(cells)}")
