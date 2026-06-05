from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.worksheet.datavalidation import DataValidation
from openpyxl.utils import get_column_letter


ROOT = Path(__file__).resolve().parents[1]

FILES = {
    "prep": ROOT / "DLV-001A-启动会会前准备与签到表.xlsx",
    "vote": ROOT / "DLV-001B-启动会会议事项表决票.xlsx",
    "script": ROOT / "DLV-002A-启动会会中解说词.docx",
    "after": ROOT / "DLV-002B-启动会会后输出物与行动项清单.xlsx",
}

INK = "1F2933"
HEADER = "E8EEF5"
SECTION = "F7F2E8"
ACCENT = "8A3B2E"
SAGE = "DDE8DF"
GOLD = "F3E4B5"
BLUE = "DDEAF6"
GRID = "C9D1D9"


def style_ws(ws, title, subtitle=None, freeze="A4"):
    ws.sheet_view.showGridLines = False
    ws.freeze_panes = freeze
    ws["A1"] = title
    ws["A1"].font = Font(name="Microsoft YaHei", size=16, bold=True, color=ACCENT)
    ws["A1"].alignment = Alignment(vertical="center")
    ws.row_dimensions[1].height = 28
    if subtitle:
        ws["A2"] = subtitle
        ws["A2"].font = Font(name="Microsoft YaHei", size=10, color="666666")
        ws["A2"].alignment = Alignment(wrap_text=True)
        ws.row_dimensions[2].height = 32


def write_table(ws, start_row, headers, rows, widths, title=None):
    if title:
        ws.cell(start_row, 1, title)
        ws.cell(start_row, 1).font = Font(name="Microsoft YaHei", size=12, bold=True, color=INK)
        ws.cell(start_row, 1).fill = PatternFill("solid", fgColor=SECTION)
        ws.merge_cells(start_row=start_row, start_column=1, end_row=start_row, end_column=len(headers))
        start_row += 1

    thin = Side(style="thin", color=GRID)
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for c, h in enumerate(headers, 1):
        cell = ws.cell(start_row, c, h)
        cell.font = Font(name="Microsoft YaHei", size=10, bold=True, color=INK)
        cell.fill = PatternFill("solid", fgColor=HEADER)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = border
        ws.column_dimensions[get_column_letter(c)].width = widths[c - 1]

    for r, row in enumerate(rows, start_row + 1):
        for c, value in enumerate(row, 1):
            cell = ws.cell(r, c, value)
            cell.font = Font(name="Microsoft YaHei", size=10, color=INK)
            cell.alignment = Alignment(vertical="center", wrap_text=True)
            cell.border = border
            if c in (1, 5, 6, 7):
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        ws.row_dimensions[r].height = 34

    ws.auto_filter.ref = f"A{start_row}:{get_column_letter(len(headers))}{start_row + len(rows)}"
    return start_row + len(rows) + 2


def add_validation(ws, cell_range, values):
    dv = DataValidation(type="list", formula1=f'"{",".join(values)}"', allow_blank=True)
    ws.add_data_validation(dv)
    dv.add(cell_range)


def save_wb(wb, path):
    for ws in wb.worksheets:
        ws.page_setup.orientation = "landscape"
        ws.page_setup.fitToWidth = 1
        ws.page_setup.fitToHeight = 0
        ws.sheet_properties.pageSetUpPr.fitToPage = True
        ws.page_margins.left = 0.35
        ws.page_margins.right = 0.35
        ws.page_margins.top = 0.5
        ws.page_margins.bottom = 0.5
    wb.save(path)


def create_prep_workbook():
    wb = Workbook()
    ws = wb.active
    ws.title = "会前准备清单"
    style_ws(ws, "启动会会前准备清单", "用于会前逐项检查材料、人员、会场、表决、证据留存等准备情况。")
    rows = [
        [1, "会议通知", "发布启动会通知，明确时间、地点、议程、参会范围", "PMO", "2026-06-07", "未开始", ""],
        [2, "会议材料", "确认启动会议程与会议事项表决清单", "刘春含 / 张琇雅", "2026-06-07", "未开始", ""],
        [3, "计划材料", "准备总体实施计划和关键里程碑材料", "刘春含", "2026-06-07", "未开始", ""],
        [4, "技术材料", "准备业务流程调研现状和工作组分工材料", "张广懿", "2026-06-07", "未开始", ""],
        [5, "授权材料", "准备 PMO 治理与调研授权说明", "李洪哲 / PMO", "2026-06-07", "未开始", ""],
        [6, "签到表", "打印签到表，预留部门代表和列席人员空白行", "张琇雅", "2026-06-08", "未开始", ""],
        [7, "表决票", "准备会议事项表决票和表决统计表", "张琇雅", "2026-06-08", "未开始", ""],
        [8, "会场", "确认投影、音响、座席、名牌、白板和网络", "PMO", "2026-06-08", "未开始", ""],
        [9, "证据留存", "准备录音录像设备，并在会前说明用途和边界", "PMO", "2026-06-08", "未开始", ""],
        [10, "启动令", "准备项目启动令模板和签发页", "PMO", "2026-06-08", "未开始", ""],
        [11, "责任池", "准备责任池事项登记表模板", "PMO", "2026-06-08", "未开始", ""],
        [12, "部门联络", "准备各部门联络人收集表", "各部门 / PMO", "2026-06-15", "未开始", ""],
    ]
    write_table(ws, 4, ["序号", "模块", "准备事项", "责任人/部门", "截止时间", "状态", "备注"], rows, [8, 14, 42, 20, 14, 12, 24])
    add_validation(ws, "F5:F40", ["未开始", "进行中", "已完成", "需协调"])

    ws2 = wb.create_sheet("签到表")
    style_ws(ws2, "启动会签到表", "实际参会人员以本表为准；签到表随会议纪要归档。")
    signin_rows = [
        [1, "马成文", "项目决策层", "开场与项目背景、总结与签发启动令", "线下", "", "", ""],
        [2, "李洪哲", "项目决策层", "PMO 治理与调研授权发布、签发启动令", "线下", "", "", ""],
        [3, "刘春含", "PMO", "主持、总体实施计划评审、关键里程碑确认、表决宣读", "线下", "", "", ""],
        [4, "张广懿", "技术解说", "业务流程调研现状、工作组分工与职责", "线下", "", "", ""],
        [5, "张琇雅", "PMO", "会议纪要、表决结果记录", "线下", "", "", ""],
        [6, "", "经营发展部", "部门代表", "线下", "", "", ""],
        [7, "", "项目管理部", "部门代表", "线下", "", "", ""],
        [8, "", "行政人事部", "部门代表", "线下", "", "", ""],
        [9, "", "物资保障部", "部门代表", "线下", "", "", ""],
        [10, "", "财务部", "部门代表", "线下", "", "", ""],
        [11, "", "运维安环部", "部门代表", "线下", "", "", ""],
    ] + [[i, "", "", "", "", "", "", ""] for i in range(12, 31)]
    write_table(ws2, 4, ["序号", "姓名", "部门/单位", "会议角色", "出席方式", "联系方式", "签名", "备注"], signin_rows, [7, 12, 18, 36, 12, 16, 18, 20])
    add_validation(ws2, "E5:E40", ["线下", "线上", "请假", "代会"])

    ws3 = wb.create_sheet("会场检查")
    style_ws(ws3, "会场与设备检查表", "会前 30 分钟完成最后检查。")
    venue_rows = [
        [1, "会场座席", "座席、名牌、签到台已摆放", "PMO", "未检查", ""],
        [2, "投影设备", "投影、翻页器、转接头可用", "PMO", "未检查", ""],
        [3, "音频设备", "麦克风、音响、录音设备可用", "PMO", "未检查", ""],
        [4, "网络接入", "线上接入或备用热点可用", "PMO", "未检查", ""],
        [5, "材料打印", "议程、签到表、表决票、启动令签发页已打印", "张琇雅", "未检查", ""],
        [6, "证据告知", "录音录像目的和使用边界已在会前说明", "主持人", "未检查", ""],
    ]
    write_table(ws3, 4, ["序号", "检查项", "检查内容", "责任人", "结果", "备注"], venue_rows, [7, 18, 44, 16, 14, 24])
    add_validation(ws3, "E5:E30", ["未检查", "通过", "需处理"])

    ws4 = wb.create_sheet("资料接收")
    style_ws(ws4, "会前资料接收表", "用于记录各发言材料和会后输出物模板是否准备到位。")
    material_rows = [
        [1, "PMO治理与调研授权说明", "李洪哲 / PMO", "启动授权", "张琇雅", "未收到", ""],
        [2, "总体实施计划", "刘春含", "计划评审", "张琇雅", "未收到", ""],
        [3, "关键里程碑清单", "刘春含", "里程碑确认", "张琇雅", "未收到", ""],
        [4, "业务流程调研现状材料", "张广懿", "流程调研", "张琇雅", "未收到", ""],
        [5, "工作组分工说明", "张广懿", "职责确认", "张琇雅", "未收到", ""],
        [6, "会议事项表决票", "张琇雅", "会议表决", "张琇雅", "未收到", ""],
        [7, "项目启动令模板", "PMO", "启动令签发", "张琇雅", "未收到", ""],
    ]
    write_table(ws4, 4, ["序号", "资料名称", "提供方", "用途", "接收人", "状态", "备注"], material_rows, [7, 30, 18, 18, 14, 12, 24])
    add_validation(ws4, "F5:F30", ["未收到", "已收到", "需补充", "已归档"])
    save_wb(wb, FILES["prep"])


def create_vote_workbook():
    wb = Workbook()
    ws = wb.active
    ws.title = "表决票"
    style_ws(ws, "启动会会议事项表决票", "由刘春含逐项宣读，全体投票；表决结果随启动会纪要归档。")
    items = [
        ["V-01", "项目启动", "同意数字化底座项目进入启动执行状态"],
        ["V-02", "PMO治理授权", "同意 PMO 对项目计划、交付物、阶段门、风险、问题、变更和责任池事项进行治理"],
        ["V-03", "PMO调研授权", "同意 PMO 开展现行流程治理、录音录像、全日跟产或参会、业务行为模拟、历史不符合项推演"],
        ["V-04", "责任池原则", "同意启动期历史问题、边界问题和资料不一致问题纳入责任池"],
        ["V-05", "不追责边界", "除拖期、交付物未交、提交材料质量过于低劣外，责任池内事项一律不追责"],
        ["V-06", "部门配合机制", "同意各部门指定联络人，配合资料提供、现场说明、流程确认和阶段性反馈"],
        ["V-07", "默认推进规则", "同意逾期未反馈事项由 PMO 形成默认版本进入下一轮评审"],
        ["V-08", "启动令签发", "同意会后按表决结果签发项目启动令"],
    ]
    rows = [[i + 1, *item, "", "", "", "", ""] for i, item in enumerate(items)]
    write_table(ws, 4, ["序号", "编号", "表决事项", "表决口径", "同意", "修改后同意", "暂缓", "不同意", "备注"], rows, [7, 10, 18, 58, 10, 14, 10, 10, 22])

    ws2 = wb.create_sheet("表决统计")
    style_ws(ws2, "会议事项表决统计", "统计可会后补录；通过口径可由项目决策层确认。")
    stat_rows = [[item[0], item[1], "", "", "", "", "", ""] for item in items]
    write_table(ws2, 4, ["编号", "表决事项", "应投票数", "同意", "修改后同意", "暂缓", "不同意", "结论"], stat_rows, [10, 20, 12, 10, 14, 10, 10, 14])
    add_validation(ws2, "H5:H20", ["通过", "未通过", "暂缓", "需复议"])

    ws3 = wb.create_sheet("投票人清单")
    style_ws(ws3, "投票人清单", "用于确认参与表决人员范围。")
    voters = [
        [1, "马成文", "项目决策层", "投票人", ""],
        [2, "李洪哲", "项目决策层", "投票人", ""],
        [3, "刘春含", "PMO", "宣读 / 投票人", ""],
        [4, "张广懿", "技术解说", "投票人", ""],
        [5, "张琇雅", "PMO", "记录人", ""],
    ] + [[i, "", "", "", ""] for i in range(6, 31)]
    write_table(ws3, 4, ["序号", "姓名", "部门/单位", "表决角色", "签名"], voters, [7, 14, 20, 18, 24])
    save_wb(wb, FILES["vote"])


def create_after_workbook():
    wb = Workbook()
    ws = wb.active
    ws.title = "会后输出物"
    style_ws(ws, "启动会会后输出物清单", "用于跟踪启动会后必须归档或下发的材料。")
    outputs = [
        ["OUT-001", "启动会纪要", "会议全过程", "张琇雅", "2026-06-08", "docx/pdf", "未开始", ""],
        ["OUT-002", "会议事项表决结果", "会议事项表决", "张琇雅", "2026-06-08", "xlsx/pdf", "未开始", ""],
        ["OUT-003", "项目启动令", "总结与签发启动令", "马成文 / 李洪哲", "2026-06-08", "docx/pdf", "未开始", ""],
        ["OUT-004", "项目治理与调研机制管理办法", "PMO治理与调研授权", "信息化项目管理工作室", "2026-06-15", "docx/md", "未开始", ""],
        ["OUT-005", "业务流程调研资料清单", "业务流程调研现状", "张广懿", "2026-06-15", "xlsx", "未开始", ""],
        ["OUT-006", "工作组职责确认表", "工作组分工与职责", "张广懿 / PMO", "2026-06-15", "xlsx/docx", "未开始", ""],
        ["OUT-007", "责任池事项登记表", "责任池原则", "PMO", "2026-06-15", "xlsx", "未开始", ""],
        ["OUT-008", "部门联络人清单", "部门配合机制", "各部门 / PMO", "2026-06-15", "xlsx", "未开始", ""],
    ]
    write_table(ws, 4, ["编号", "输出物", "来源议程/事项", "责任人", "截止时间", "交付形态", "状态", "备注"], outputs, [12, 30, 24, 22, 14, 14, 12, 24])
    add_validation(ws, "G5:G40", ["未开始", "进行中", "已提交", "已归档", "需协调"])

    ws2 = wb.create_sheet("行动项")
    style_ws(ws2, "会后行动项清单", "用于 PMO 周会跟踪，必要时升级项目决策层。")
    actions = [
        ["A-01", "发布启动会纪要", "张琇雅", "PMO", "2026-06-08", "高", "未开始", "纪要归档"],
        ["A-02", "整理会议事项表决结果", "张琇雅", "刘春含", "2026-06-08", "高", "未开始", "表决结果记录"],
        ["A-03", "签发项目启动令", "马成文 / 李洪哲", "PMO", "2026-06-08", "高", "未开始", "启动令"],
        ["A-04", "发布项目治理与调研机制管理办法", "信息化项目管理工作室", "各部门", "2026-06-15", "高", "未开始", "DLV-003"],
        ["A-05", "形成业务流程调研资料清单", "张广懿", "各部门", "2026-06-15", "中", "未开始", "资料清单"],
        ["A-06", "建立责任池事项登记表", "PMO", "各部门", "2026-06-15", "中", "未开始", "责任池台账"],
        ["A-07", "确认各部门联络人", "各部门", "PMO", "2026-06-15", "中", "未开始", "联络人清单"],
    ]
    write_table(ws2, 4, ["编号", "行动项", "责任方", "配合方", "截止时间", "优先级", "状态", "关闭依据"], actions, [10, 34, 22, 18, 14, 10, 12, 24])
    add_validation(ws2, "F5:F40", ["高", "中", "低"])
    add_validation(ws2, "G5:G40", ["未开始", "进行中", "已完成", "需协调", "已升级"])

    ws3 = wb.create_sheet("责任池登记")
    style_ws(ws3, "责任池事项登记表", "启动期历史问题先治理不追责；拖期、交付物未交、材料质量过于低劣除外。")
    rp_rows = [["RP-001", "", "", "", "", "", "", "否", "", "待核实"]] + [["", "", "", "", "", "", "", "", "", ""] for _ in range(20)]
    write_table(ws3, 4, ["编号", "发现日期", "事项名称", "类型", "涉及部门", "事实描述", "证据来源", "是否追责例外", "治理动作", "状态"], rp_rows, [10, 14, 24, 18, 18, 34, 24, 16, 30, 14])
    add_validation(ws3, "D5:D40", ["历史流程问题", "历史数据问题", "边界问题", "表单台账不一致", "历史不符合项", "执行问题"])
    add_validation(ws3, "H5:H40", ["是", "否"])
    add_validation(ws3, "J5:J40", ["待核实", "治理中", "已关闭", "转问题台账"])

    ws4 = wb.create_sheet("资料补交清单")
    style_ws(ws4, "资料补交清单", "用于会后向部门收集调研资料。")
    material_rows = [
        ["M-001", "制度文件与流程说明", "各部门", "现行流程治理", "2026-06-15", "张琇雅", "未开始", ""],
        ["M-002", "表单、台账、会议记录样例", "各部门", "流程地图和数据地图", "2026-06-15", "张琇雅", "未开始", ""],
        ["M-003", "典型业务样例单据", "各部门", "业务行为模拟", "2026-06-20", "张广懿", "未开始", ""],
        ["M-004", "历史不符合项或历史问题材料", "相关部门", "历史不符合项推演", "2026-06-20", "PMO", "未开始", ""],
    ]
    write_table(ws4, 4, ["编号", "资料名称", "提供方", "用途", "截止时间", "接收人", "状态", "备注"], material_rows, [10, 32, 18, 24, 14, 14, 12, 24])
    add_validation(ws4, "G5:G40", ["未开始", "已通知", "已收到", "需补充", "已归档"])
    save_wb(wb, FILES["after"])


def set_font(run, name="Microsoft YaHei", size=10.5, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_para_style(p, before=0, after=6, line=1.15):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_para_style(p, before=14 if level == 1 else 8, after=5)
    run = p.add_run(text)
    set_font(run, size=15 if level == 1 else 12.5, color=ACCENT if level == 1 else "1F4D78", bold=True)
    return p


def add_body(doc, text, bold_label=None):
    p = doc.add_paragraph()
    set_para_style(p, after=5)
    if bold_label:
        r = p.add_run(bold_label)
        set_font(r, bold=True)
    r = p.add_run(text)
    set_font(r)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc, headers, rows, widths_cm):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        shade_cell(hdr[i], HEADER)
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                set_font(r, size=9.5, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                set_para_style(p, after=0, line=1.1)
                for r in p.runs:
                    set_font(r, size=9)
    for row in table.rows:
        for i, width in enumerate(widths_cm):
            row.cells[i].width = Cm(width)
    return table


def create_script_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("昌兴复材数字化底座项目启动会会中解说词")
    set_font(r, size=20, color=ACCENT, bold=True)
    set_para_style(title, after=4)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("主持人：刘春含｜技术解说：张广懿｜会议纪要：张琇雅")
    set_font(r, size=10.5, color="555555")
    set_para_style(subtitle, after=14)

    add_table(
        doc,
        ["议程", "主讲 / 执行", "目标"],
        [
            ["开场与项目背景", "马成文", "统一项目背景和启动要求"],
            ["PMO治理与调研授权发布", "李洪哲", "明确 PMO 授权、调研方式和责任池原则"],
            ["总体实施计划评审", "刘春含", "确认总体节奏和 WBS1 启动安排"],
            ["关键里程碑确认", "刘春含", "确认阶段门和会后行动项"],
            ["业务流程调研现状", "张广懿", "说明现状、缺口和调研方向"],
            ["工作组分工与职责", "张广懿", "明确工作组职责和部门配合边界"],
            ["会议事项表决", "刘春含宣读、全体投票", "形成表决结果"],
            ["总结与签发启动令", "马成文、李洪哲", "确认启动令签发和执行要求"],
        ],
        [4.2, 4.0, 7.0],
    )

    add_heading(doc, "一、主持开场口径", 1)
    add_body(doc, "各位领导、各位同事，大家上午好。今天召开昌兴复材数字化底座项目启动会。本次会议由我主持，张广懿负责技术解说，张琇雅负责会议纪要和会后材料整理。", "刘春含：")
    add_body(doc, "本次会议的重点不是简单宣布一个信息化项目开始，而是确认项目治理规则、调研方式、责任池边界和后续工作分工。会议将按照既定议程推进，会议事项由我逐项宣读，全体投票形成表决结果。")

    sections = [
        (
            "二、开场与项目背景",
            "马成文",
            [
                "本项目的启动，核心是为了把公司流程、数据、责任边界和后续数字化建设基础梳理清楚。当前阶段不是先评价哪个应用系统，而是先把真实流程、真实数据、真实问题和真实交接关系弄清楚。",
                "各部门要把本次项目理解为一次共同治理。项目推进过程中，历史遗留问题先进入责任池治理，不直接追究部门或个人责任。我们要鼓励真实暴露问题，而不是让问题继续藏在纸面流程和口头习惯里。",
                "请 PMO 按照项目统一节奏组织推进，各部门按会议确认的机制配合资料、调研、现场说明和阶段性确认。",
            ],
        ),
        (
            "三、PMO治理与调研授权发布",
            "李洪哲",
            [
                "会议确认后，PMO 将作为项目治理牵头单位，对项目计划、交付物、阶段门、风险、问题、变更和责任池事项进行统一管理。",
                "PMO 对各部门开展调研时，可以采用现行流程治理、录音录像、全日跟产或参会、业务行为模拟、历史不符合项推演等方式。相关记录用于流程地图、数据地图、台账闭环和阶段门评审。",
                "责任池是本次启动期的重要机制。启动、调研、流程地图、数据地图和现状盘点阶段暴露的历史流程问题、历史数据问题、跨部门边界问题、表单台账不一致问题和历史不符合项，原则上先治理，不追责。",
                "但责任池不是免做池。拖期、交付物未交、提交材料质量过于低劣这三类执行问题，不纳入责任池保护。",
            ],
        ),
        (
            "四、总体实施计划评审",
            "刘春含",
            [
                "总体实施计划按照 WBS 管理。WBS1 是项目启动与总体蓝图，当前重点是启动会准备、启动会召开、项目治理机制建立，以及后续现状调研计划和业务流程调研。",
                "后续工作不是靠口头推进，而是通过会议纪要、行动项、交付物、台账和阶段性确认形成闭环。凡是影响阶段门、交付物、关键路径、部门配合和后续调研的事项，都要进入 PMO 管控视图。",
                "请各工作组按照 WBS 节奏推进；未反馈事项将由 PMO 形成默认版本进入下一轮评审，并在周会中披露阻塞事实和影响。",
            ],
        ),
        (
            "五、关键里程碑确认",
            "刘春含",
            [
                "启动会后，第一项里程碑是启动会纪要和会议事项表决结果归档。第二项是项目启动令签发。第三项是项目治理与调研机制管理办法发布。",
                "随后进入现状调研计划编制，调研计划要明确调研对象、调研方式、资料清单、责任人、时间安排、输出物和责任池口径。",
                "阶段性确认不是无限背责。部门确认口径是：该版本可作为下一阶段调研、设计、开发、测试或联调依据；后续变化按变更流程处理。",
            ],
        ),
        (
            "六、业务流程调研现状",
            "张广懿",
            [
                "目前公司已有大量制度、程序文件、表单、台账和实际业务经验，但这些材料之间不一定完全一致。流程地图和数据地图要解决的，就是制度怎么写、现场怎么做、表单怎么走、数据在哪里产生、责任在哪里交接。",
                "调研时不会要求部门临时创造一套漂亮流程。PMO 和工作组要看真实业务：真实单据、真实会议、真实等待、真实返工、真实例外处理。",
                "对于历史不符合项，我们会用推演方式倒推流程断点和数据断点。上线前暴露的历史问题进入责任池治理；这有利于后续系统设计、数据治理和阶段门评审。",
            ],
        ),
        (
            "七、工作组分工与职责",
            "张广懿",
            [
                "各工作组按照专业边界承担不同输入。流程与数据相关工作组负责调研组织、流程梳理、数据对象识别和问题记录；技术架构相关工作组负责把调研结果转化为后续方案约束；基础设施和安全相关工作组负责支撑环境、权限和访问边界。",
                "业务部门负责提供制度、表单、台账、样例、现场说明和阶段性确认。部门联络人负责日常沟通，但部门级决策和阶段性确认仍由相应负责人确认。",
                "遇到跨部门边界不清、关键材料缺失、阶段门争议、供应商责任边界争议时，按 PMO 周会和项目决策层路径升级。",
            ],
        ),
        (
            "八、会议事项表决",
            "刘春含宣读、全体投票",
            [
                "下面进入会议事项表决环节。请各位按表决票逐项确认。表决事项包括：项目启动、PMO 治理授权、PMO 调研授权、责任池原则、不追责边界、部门配合机制、默认推进规则、启动令签发。",
                "请注意，责任池保护的是启动期暴露历史问题的积极性。除拖期、交付物未交、提交材料质量过于低劣外，责任池内事项一律不追责。",
                "表决结果由张琇雅记录，会后随会议纪要一起归档。",
            ],
        ),
        (
            "九、总结与签发启动令",
            "马成文、李洪哲",
            [
                "本次会议的核心结论是：项目正式启动，PMO 治理与调研授权生效，责任池机制生效，各部门按会议确认的方式配合后续调研和交付。",
                "会后由 PMO 整理会议纪要、表决结果和行动项，由相关负责人签发项目启动令。请各部门把这次项目当作流程和数据治理的共同工程，而不是单一系统建设任务。",
                "启动令签发后，PMO 按 WBS 和阶段门组织推进，各工作组按职责执行，各部门按资料清单和调研安排配合。",
            ],
        ),
    ]

    for heading, speaker, lines in sections:
        add_heading(doc, heading, 1)
        for idx, line in enumerate(lines):
            label = f"{speaker}：" if idx == 0 else None
            add_body(doc, line, label)

    add_heading(doc, "十、现场问答备用口径", 1)
    backup = [
        ["如果部门担心追责", "本项目启动期设置责任池，历史问题先治理不追责。只有拖期、交付物未交、提交材料质量过于低劣这三类执行问题不受责任池保护。"],
        ["如果有人问为什么要录音录像", "录音录像用于事实还原、纪要校核和争议复核，不用于个人绩效评价，也不脱离项目目的传播。"],
        ["如果有人认为资料不完整", "资料不完整本身可以进入责任池治理。请提供现有版本和差异说明，PMO 会按影响范围和后续用途分级处理。"],
        ["如果部门无法按期反馈", "请提前说明原因并提出可行时间。逾期未反馈且未说明的，PMO 可形成默认版本进入下一轮评审。"],
    ]
    add_table(doc, ["场景", "建议回答"], backup, [4.2, 11.0])

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("昌兴复材数字化底座项目启动会｜会中解说词")
    set_font(r, size=9, color="777777")

    doc.save(FILES["script"])


if __name__ == "__main__":
    create_prep_workbook()
    create_vote_workbook()
    create_script_docx()
    create_after_workbook()
    for key, path in FILES.items():
        print(f"{key}: {path}")
