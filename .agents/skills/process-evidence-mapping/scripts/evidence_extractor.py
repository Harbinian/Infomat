#!/usr/bin/env python
"""Extract traceable evidence chunks for process-evidence vector retrieval."""

from __future__ import annotations

import argparse
import csv
import hashlib
import html
import json
import os
import re
import shutil
import sys
import tempfile
from pathlib import Path
from xml.etree import ElementTree


TEXT_EXTENSIONS = {".md", ".txt", ".csv", ".json", ".html", ".htm"}
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | {".docx", ".doc", ".xlsx", ".xls", ".pdf", ".vsd", ".vsdx"}
SKIP_DIRS = {".git", "node_modules", "artifacts", "test-results", "__pycache__", "_extracted", "流程治理"}
GENERATED_FILE_PATTERNS = [
    re.compile(r".*部门-能力-流程-系统映射关系\.md$"),
    re.compile(r".*能力层与MDM建设要求\.md$"),
    re.compile(r".*部门能力流程系统桑基图\.html$"),
    re.compile(r"_quality-report\.md$"),
    re.compile(r"部门能力流程系统映射使用说明\.md$"),
    re.compile(r"桑基图反馈操作说明\.html$"),
    re.compile(r"README\.md$"),
    re.compile(r"CLAUDE\.md$"),
]


def sha1(text: str) -> str:
    return hashlib.sha1(text.encode("utf-8", errors="ignore")).hexdigest()


def parse_exts(value: str | None) -> set[str]:
    if not value:
        return set()
    exts: set[str] = set()
    for item in re.split(r"[,;\s]+", value):
        item = item.strip().lower()
        if not item:
            continue
        exts.add(item if item.startswith(".") else f".{item}")
    return exts


def repo_path(path: Path, repo: Path) -> str:
    try:
        return path.resolve().relative_to(repo.resolve()).as_posix()
    except ValueError:
        return str(path.resolve()).replace("\\", "/")


def normalize_text(text: str) -> str:
    text = html.unescape(text)
    text = re.sub(r"<script[\s\S]*?</script>", " ", text, flags=re.I)
    text = re.sub(r"<style[\s\S]*?</style>", " ", text, flags=re.I)
    text = re.sub(r"<[^>]+>", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def detect_doc_no(file_name: str) -> tuple[str, str]:
    stem = Path(file_name).stem
    match = re.match(r"^([A-Za-z0-9]+(?:-[A-Za-z0-9]+)+)", stem)
    if not match:
        return "", ""
    doc_no = match.group(1)
    version_match = re.search(r"-([A-Z])$", doc_no)
    return doc_no, version_match.group(1) if version_match else ""


def clause_from_text(text: str) -> str:
    match = re.match(r"^\s*(\d+(?:\.\d+){0,5})\s*[\.、\s]", text)
    return match.group(1) if match else ""


def looks_heading(text: str) -> bool:
    return bool(clause_from_text(text)) or text.strip() in {
        "前言",
        "范围",
        "规范性引用文件",
        "术语和定义",
        "职责",
        "工作程序",
        "记录",
        "附录",
    }


def normalized_candidate(raw: str) -> str:
    candidates: list[str] = []
    if re.search(r"公司\s+月综合打分表", raw) or "公司_月综合打分表" in raw or "公司__月综合打分表" in raw:
        candidates.append("公司__月综合打分表")
        candidates.append("公司月度综合打分表候选")
    if "__" in raw:
        candidates.append(raw.replace("__", "_"))
    if re.search(r"[\u4e00-\u9fff]\s{2,}[\u4e00-\u9fff]", raw):
        candidates.append(re.sub(r"\s{2,}", "", raw))
    return " / ".join(dict.fromkeys(candidate for candidate in candidates if candidate))


def extraction_quality(raw: str, status: str = "clean") -> str:
    if status in {"failed", "needs_ocr"}:
        return status
    if normalized_candidate(raw):
        return "partial"
    if "__" in raw or re.search(r"[\u4e00-\u9fff]\s+月综合打分表", raw):
        return "partial"
    return "clean"


def base_source(file_path: Path, repo: Path, extraction_status: str) -> dict:
    stat = file_path.stat()
    doc_no, version = detect_doc_no(file_path.name)
    rel = repo_path(file_path, repo)
    return {
        "source_file_id": sha1(rel)[:16],
        "source_file": rel,
        "source_file_name": file_path.name,
        "leaf_dir": repo_path(file_path.parent, repo),
        "doc_no": doc_no,
        "version": version,
        "source_company": "",
        "source_org_name": "",
        "file_ext": file_path.suffix.lower(),
        "file_size": stat.st_size,
        "modified_time": stat.st_mtime,
        "extraction_status": extraction_status,
        "included_status": "candidate",
        "included_reason": "Chunked for retrieval review only; inclusion still requires source verification.",
    }


def chunk_record(source: dict, raw: str, artifact_type: str, index: str, **extra: str) -> dict:
    raw = normalize_text(raw)
    quality = extraction_quality(raw)
    candidate = normalized_candidate(raw)
    return {
        "chunk_id": f"{source['source_file_id']}-{index}",
        "source_file_id": source["source_file_id"],
        "source_file": source["source_file"],
        "source_file_name": source["source_file_name"],
        "leaf_dir": source["leaf_dir"],
        "doc_no": source.get("doc_no", ""),
        "version": source.get("version", ""),
        "source_company": source.get("source_company", ""),
        "source_org_name": source.get("source_org_name", ""),
        "clause": extra.get("clause", ""),
        "clause_title": extra.get("clause_title", ""),
        "paragraph_id": extra.get("paragraph_id", ""),
        "table_id": extra.get("table_id", ""),
        "row_id": extra.get("row_id", ""),
        "column_name": extra.get("column_name", ""),
        "sheet_name": extra.get("sheet_name", ""),
        "form_name": extra.get("form_name", ""),
        "raw_text": raw,
        "normalized_text": raw,
        "normalized_candidate": candidate,
        "artifact_type": artifact_type,
        "extraction_method": extra.get("extraction_method", source.get("extraction_method", "python")),
        "extraction_quality": quality,
        "retrieval_method": "chunking",
        "evidence_status": "candidate",
        "verification_status": "unverified",
        "review_required": True,
        "review_reason": "Retrieval chunk only; verify original source before using in mapping.",
        "allowed_downstream_use": "review_only",
        "chunk_hash": sha1(raw),
    }


def form_name_from_text(text: str) -> str:
    if "公司月度综合打分表" in text or "公司__月综合打分表" in text or "公司 月综合打分表" in text:
        return "公司月度综合打分表"
    if "工作任务调整申请单" in text:
        return "工作任务调整申请单"
    if "经营发展部绩效评分表" in text:
        return "经营发展部绩效评分表"
    return ""


def extract_text_file(file_path: Path, repo: Path) -> tuple[dict, list[dict], list[str]]:
    source = base_source(file_path, repo, "chunked")
    text = file_path.read_text(encoding="utf-8", errors="ignore")
    chunks: list[dict] = []
    section_title = ""
    paragraph_id = 0
    table_id = 0
    row_id = 0
    buffer: list[str] = []

    def flush() -> None:
        nonlocal paragraph_id, buffer
        raw = "\n".join(buffer).strip()
        if not raw:
            return
        paragraph_id += 1
        chunks.append(chunk_record(
            source,
            raw,
            "body",
            f"P{paragraph_id:04d}",
            paragraph_id=f"P{paragraph_id}",
            clause=clause_from_text(section_title),
            clause_title=section_title,
            extraction_method="text",
        ))
        buffer = []

    for line in text.replace("\r\n", "\n").split("\n"):
        trimmed = line.strip()
        if re.match(r"^#{1,6}\s+", trimmed):
            flush()
            section_title = re.sub(r"^#{1,6}\s+", "", trimmed).strip()
            row_id = 0
            continue
        if trimmed.startswith("|") and trimmed.endswith("|"):
            flush()
            if row_id == 0:
                table_id += 1
            row_id += 1
            if not re.match(r"^\|\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?$", trimmed):
                chunks.append(chunk_record(
                    source,
                    trimmed,
                    "table",
                    f"T{table_id:02d}R{row_id:03d}",
                    table_id=f"T{table_id:02d}",
                    row_id=str(row_id),
                    clause=clause_from_text(section_title),
                    clause_title=section_title,
                    extraction_method="text",
                ))
            continue
        if not trimmed:
            flush()
            row_id = 0
            continue
        buffer.append(line)
    flush()
    return source, chunks, []


def extract_docx(file_path: Path, repo: Path) -> tuple[dict, list[dict], list[str]]:
    from docx import Document

    source = base_source(file_path, repo, "chunked")
    source["extraction_method"] = "python-docx"
    doc = Document(str(file_path))
    chunks: list[dict] = []
    current_clause = ""
    current_title = ""
    paragraph_count = 0

    for idx, paragraph in enumerate(doc.paragraphs, start=1):
        raw = normalize_text(paragraph.text)
        if not raw:
            continue
        clause = clause_from_text(raw)
        if clause or looks_heading(raw):
            if clause:
                current_clause = clause
            current_title = raw[:120]
        paragraph_count += 1
        chunks.append(chunk_record(
            source,
            raw,
            "body",
            f"P{paragraph_count:04d}",
            paragraph_id=f"P{idx}",
            clause=current_clause,
            clause_title=current_title,
            form_name=form_name_from_text(raw),
            extraction_method="python-docx",
        ))

    for table_index, table in enumerate(doc.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            cells: list[str] = []
            for cell in row.cells:
                value = normalize_text(cell.text)
                if value and (not cells or cells[-1] != value):
                    cells.append(value)
            if not cells:
                continue
            raw = f"Table T{table_index:02d} Row {row_index:03d}: {' | '.join(cells)}"
            chunks.append(chunk_record(
                source,
                raw,
                "table",
                f"T{table_index:02d}R{row_index:03d}",
                table_id=f"T{table_index:02d}",
                row_id=str(row_index),
                form_name=form_name_from_text(raw),
                extraction_method="python-docx",
            ))
    return source, chunks, []


def convert_doc_to_docx(file_path: Path, temp_dir: Path) -> Path:
    import win32com.client

    temp_input = temp_dir / f"{sha1(str(file_path))[:12]}{file_path.suffix.lower()}"
    shutil.copy2(file_path, temp_input)
    out = temp_dir / f"{sha1(str(file_path))[:12]}.docx"
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        try:
            word.AutomationSecurity = 3
        except Exception:
            pass
        doc = word.Documents.Open(str(temp_input), ReadOnly=True, AddToRecentFiles=False, ConfirmConversions=False, NoEncodingDialog=True)
        doc.SaveAs2(str(out), FileFormat=16)
        doc.Close(False)
    finally:
        word.Quit()
    return out


def extract_xlsx(file_path: Path, repo: Path) -> tuple[dict, list[dict], list[str]]:
    import openpyxl

    source = base_source(file_path, repo, "chunked")
    source["extraction_method"] = "openpyxl"
    wb = openpyxl.load_workbook(str(file_path), read_only=True, data_only=True)
    chunks: list[dict] = []
    for sheet in wb.worksheets:
        for row_index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            cells = [normalize_text(str(value)) for value in row if value is not None and normalize_text(str(value))]
            if not cells:
                continue
            raw = f"Sheet {sheet.title} Row {row_index:03d}: {' | '.join(cells)}"
            chunks.append(chunk_record(
                source,
                raw,
                "table",
                f"S{sha1(sheet.title)[:6]}R{row_index:03d}",
                sheet_name=sheet.title,
                row_id=str(row_index),
                form_name=form_name_from_text(raw),
                extraction_method="openpyxl",
            ))
    return source, chunks, []


def extract_xls(file_path: Path, repo: Path) -> tuple[dict, list[dict], list[str]]:
    import xlrd

    source = base_source(file_path, repo, "chunked")
    source["extraction_method"] = "xlrd"
    book = xlrd.open_workbook(str(file_path), on_demand=True)
    chunks: list[dict] = []
    for sheet_name in book.sheet_names():
        sheet = book.sheet_by_name(sheet_name)
        for row_index in range(sheet.nrows):
            cells = [normalize_text(str(sheet.cell_value(row_index, col))) for col in range(sheet.ncols)]
            cells = [cell for cell in cells if cell]
            if not cells:
                continue
            raw = f"Sheet {sheet_name} Row {row_index + 1:03d}: {' | '.join(cells)}"
            chunks.append(chunk_record(
                source,
                raw,
                "table",
                f"S{sha1(sheet_name)[:6]}R{row_index + 1:03d}",
                sheet_name=sheet_name,
                row_id=str(row_index + 1),
                form_name=form_name_from_text(raw),
                extraction_method="xlrd",
            ))
    return source, chunks, []


def extract_pdf(file_path: Path, repo: Path) -> tuple[dict, list[dict], list[str]]:
    source = base_source(file_path, repo, "chunked")
    chunks: list[dict] = []
    warnings: list[str] = []
    try:
        import pdfplumber
        with pdfplumber.open(str(file_path)) as pdf:
            total_text = 0
            total_images = 0
            for page_index, page in enumerate(pdf.pages, start=1):
                text = normalize_text(page.extract_text() or "")
                total_text += len(text)
                total_images += len(page.images)
                if text:
                    chunks.append(chunk_record(
                        source,
                        f"PDF Page {page_index}: {text}",
                        "body",
                        f"PG{page_index:04d}",
                        paragraph_id=f"page-{page_index}",
                        extraction_method="pdfplumber",
                    ))
            if total_text == 0 and total_images > 0:
                source["extraction_status"] = "needs_ocr"
                raw = f"待OCR/人工目视: {source['source_file']} pages={len(pdf.pages)} image_pages={total_images}"
                chunk = chunk_record(source, raw, "ocr", "OCR0001", extraction_method="pdfplumber")
                chunk["extraction_quality"] = "needs_ocr"
                chunks.append(chunk)
                warnings.append(f"needs_ocr: {source['source_file']}")
    except Exception as error:
        source["extraction_status"] = "failed"
        warnings.append(f"failed pdf: {source['source_file']} - {error}")
    return source, chunks, warnings


def extract_vsd(file_path: Path, repo: Path, temp_dir: Path) -> tuple[dict, list[dict], list[str]]:
    source = base_source(file_path, repo, "chunked")
    chunks: list[dict] = []
    warnings: list[str] = []
    try:
        import win32com.client

        temp_input = temp_dir / f"{sha1(str(file_path))[:12]}{file_path.suffix.lower()}"
        shutil.copy2(file_path, temp_input)
        visio = win32com.client.DispatchEx("Visio.Application")
        visio.Visible = False
        try:
            doc = visio.Documents.OpenEx(str(temp_input), 64)
            for page_index in range(1, int(doc.Pages.Count) + 1):
                svg = temp_dir / f"{sha1(str(file_path))[:12]}_p{page_index}.svg"
                doc.Pages.Item(page_index).Export(str(svg))
                try:
                    root = ElementTree.parse(str(svg)).getroot()
                    texts = []
                    for elem in root.iter():
                        if elem.tag.endswith("text"):
                            value = normalize_text("".join(elem.itertext()))
                            if value:
                                texts.append(value)
                    raw = f"Visio Page {page_index}: {' | '.join(texts)}"
                    if normalize_text(raw):
                        chunks.append(chunk_record(
                            source,
                            raw,
                            "flow",
                            f"VSDP{page_index:04d}",
                            paragraph_id=f"page-{page_index}",
                            extraction_method="visio-svg",
                        ))
                except Exception as parse_error:
                    warnings.append(f"failed svg parse: {source['source_file']} page {page_index} - {parse_error}")
            doc.Close()
        finally:
            visio.Quit()
    except Exception as error:
        source["extraction_status"] = "failed"
        warnings.append(f"failed vsd: {source['source_file']} - {error}")
    return source, chunks, warnings


def should_collect(path: Path, include_exts: set[str], exclude_exts: set[str]) -> bool:
    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        return False
    if any(pattern.match(path.name) for pattern in GENERATED_FILE_PATTERNS):
        return False
    if include_exts and ext not in include_exts:
        return False
    if exclude_exts and ext in exclude_exts:
        return False
    return True


def collect_files(input_path: Path, include_exts: set[str], exclude_exts: set[str]) -> list[Path]:
    input_path = input_path.resolve()
    if input_path.is_file():
        return [input_path] if should_collect(input_path, include_exts, exclude_exts) else []
    files: list[Path] = []
    for root, dirs, names in os.walk(input_path):
        dirs[:] = [name for name in dirs if name not in SKIP_DIRS]
        for name in names:
            path = Path(root) / name
            if should_collect(path, include_exts, exclude_exts):
                files.append(path)
    return sorted(files)


def defer_source(file_path: Path, repo: Path, reason: str) -> tuple[dict, list[dict], list[str]]:
    source = base_source(file_path, repo, "deferred")
    source["included_status"] = "deferred"
    source["included_reason"] = reason
    source["content_hash"] = ""
    source["chunks"] = 0
    return source, [], [f"deferred: {source['source_file']} - {reason}"]


def extract_file(file_path: Path, repo: Path, temp_dir: Path) -> tuple[dict, list[dict], list[str]]:
    ext = file_path.suffix.lower()
    try:
        if ext in TEXT_EXTENSIONS:
            return extract_text_file(file_path, repo)
        if ext == ".docx":
            return extract_docx(file_path, repo)
        if ext == ".doc":
            converted = convert_doc_to_docx(file_path, temp_dir)
            source, chunks, warnings = extract_docx(converted, repo)
            original_source = base_source(file_path, repo, "chunked")
            for key in ["source_file_id", "source_file", "source_file_name", "leaf_dir", "doc_no", "version", "file_ext", "file_size", "modified_time"]:
                source[key] = original_source[key]
            for chunk in chunks:
                chunk.update({
                    "source_file_id": source["source_file_id"],
                    "source_file": source["source_file"],
                    "source_file_name": source["source_file_name"],
                    "leaf_dir": source["leaf_dir"],
                    "doc_no": source["doc_no"],
                    "version": source["version"],
                    "extraction_method": "word-com-to-docx+python-docx",
                })
                chunk["chunk_id"] = f"{source['source_file_id']}-{chunk['chunk_id'].split('-')[-1]}"
            return source, chunks, warnings
        if ext == ".xlsx":
            return extract_xlsx(file_path, repo)
        if ext == ".xls":
            return extract_xls(file_path, repo)
        if ext == ".pdf":
            return extract_pdf(file_path, repo)
        if ext in {".vsd", ".vsdx"}:
            return extract_vsd(file_path, repo, temp_dir)
    except Exception as error:
        source = base_source(file_path, repo, "failed")
        warning = f"failed: {source['source_file']} - {type(error).__name__}: {error}"
        return source, [], [warning]
    source = base_source(file_path, repo, "unsupported")
    return source, [], [f"unsupported: {source['source_file']}"]


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True)
    parser.add_argument("--out", required=True)
    parser.add_argument("--source-index")
    parser.add_argument("--warnings")
    parser.add_argument("--temp-dir")
    parser.add_argument("--include-ext", default="")
    parser.add_argument("--exclude-ext", default="")
    parser.add_argument("--defer-ext", default="")
    parser.add_argument("--defer-reason", default="Deferred by --defer-ext for a separate extraction batch.")
    args = parser.parse_args()

    repo = Path.cwd()
    input_path = Path(args.input).resolve()
    out = Path(args.out).resolve()
    source_index = Path(args.source_index).resolve() if args.source_index else out.with_name("source_index.jsonl")
    warnings_path = Path(args.warnings).resolve() if args.warnings else out.with_name("chunking_warnings.md")
    temp_dir = Path(args.temp_dir).resolve() if args.temp_dir else out.parent / "_tmp_conversions"
    temp_dir.mkdir(parents=True, exist_ok=True)
    include_exts = parse_exts(args.include_ext)
    exclude_exts = parse_exts(args.exclude_ext)
    defer_exts = parse_exts(args.defer_ext)

    all_chunks: list[dict] = []
    sources: list[dict] = []
    warnings: list[str] = []
    for file_path in collect_files(input_path, include_exts, exclude_exts):
        if file_path.suffix.lower() in defer_exts:
            source, chunks, file_warnings = defer_source(
                file_path,
                repo,
                args.defer_reason,
            )
        else:
            source, chunks, file_warnings = extract_file(file_path, repo, temp_dir)
            content_hash = sha1("\n".join(chunk["raw_text"] for chunk in chunks))
            source["content_hash"] = content_hash
            source["chunks"] = len(chunks)
            for chunk in chunks:
                chunk["content_hash"] = content_hash
        sources.append(source)
        all_chunks.extend(chunks)
        warnings.extend(file_warnings)

    out.parent.mkdir(parents=True, exist_ok=True)
    source_index.parent.mkdir(parents=True, exist_ok=True)
    warnings_path.parent.mkdir(parents=True, exist_ok=True)
    out.write_text("".join(json.dumps(chunk, ensure_ascii=False) + "\n" for chunk in all_chunks), encoding="utf-8")
    source_index.write_text("".join(json.dumps(source, ensure_ascii=False) + "\n" for source in sources), encoding="utf-8")
    warnings_path.write_text(
        "\n".join([
            "# Chunking Warnings",
            "",
            f"- sources: {len(sources)}",
            f"- chunks: {len(all_chunks)}",
            f"- warnings: {len(warnings)}",
            "",
            *[f"- {warning}" for warning in warnings],
            "",
        ]),
        encoding="utf-8",
    )
    print(f"chunks={len(all_chunks)} sources={len(sources)} warnings={len(warnings)}", file=sys.stderr)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
