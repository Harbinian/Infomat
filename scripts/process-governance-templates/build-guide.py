"""Build the company-level Word guide for department process templates."""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = "2E74B5"
DARK_BLUE = "1F4E79"
BROWN = "704B3A"
BEIGE = "F4EDE3"
LIGHT_BLUE = "E8EEF5"
LIGHT_YELLOW = "FFF7D6"
LIGHT_GREY = "E9EEF1"
SAGE = "DDE8DF"
WHITE = "FFFFFF"
INK = "2F2A27"
MUTED = "706A65"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--data", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--asset-dir", required=True)
    return parser.parse_args()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = int(sum(widths) * 1440)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(grid_col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def style_table(table, header: bool = True) -> None:
    for row_index, row in enumerate(table.rows):
        prevent_row_split(row)
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.08
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = RGBColor.from_string(INK)
        if header and row_index == 0:
            for cell in row.cells:
                set_cell_shading(cell, BLUE)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor.from_string(WHITE)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float]):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cells[index].text = str(value)
    set_table_widths(table, widths)
    style_table(table)
    set_repeat_table_header(table.rows[0])
    return table


def add_label(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(BROWN)


def add_callout(document: Document, title: str, body: str, fill: str = BEIGE) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_widths(table, [6.5])
    prevent_row_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    title_run = paragraph.add_run(f"{title}：")
    title_run.bold = True
    title_run.font.color.rgb = RGBColor.from_string(BROWN)
    paragraph.add_run(body)


def add_bullet(document: Document, text: str, level: int = 0) -> None:
    style = "List Bullet" if level == 0 else "List Bullet 2"
    paragraph = document.add_paragraph(style=style)
    paragraph.add_run(text)


def add_number(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Number")
    paragraph.add_run(text)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.2

    for name, size, color, before, after in (
        ("Title", 24, BROWN, 0, 12),
        ("Subtitle", 12, MUTED, 0, 12),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, DARK_BLUE, 14, 7),
        ("Heading 3", 11.5, BROWN, 10, 5),
    ):
        style = document.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_header_footer(document: Document, snapshot_date: str) -> None:
    for section in document.sections:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        p.text = "流程与数据梳理｜部门填报与评审标准"
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in p.runs:
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string(MUTED)
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"流程映射快照 {snapshot_date}  ｜  ")
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = "PAGE"
        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        run = p.add_run()
        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_end)
        for run in p.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string(MUTED)


def find_example(data: dict) -> tuple[dict, list[dict]]:
    for department in data["departments"]:
        for process in department["processes"]:
            if process["processId"] == "FIN-L3-001":
                behaviors = [item for item in department["behaviors"] if item["processId"] == process["processId"]]
                return process, behaviors
    raise ValueError("Expected process FIN-L3-001 was not found in normalized data")


def build_document(data: dict, output: Path, asset_dir: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    configure_styles(document)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("流程与数据梳理\n填写及评审标准")
    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("九部门流程与数据字典模板配套说明｜原文制度追溯版")

    summary = document.add_table(rows=4, cols=2)
    summary.style = "Table Grid"
    for row, values in enumerate(
        [
            ["适用对象", "工程技术部、质量管理部、财务部、行政人事部、经营发展部、物资保障部、项目管理部、复材车间、运维安环部"],
            ["填报真源", "各部门 Excel 工作簿；Word 仅解释口径，不重复维护流程数据"],
            ["数据范围", f"{data['totals']['processes']} 条 L3、{data['totals']['behaviors']} 条 A1；其中 {data['totals']['unmappedProcesses']} 条系统承接方向待确认"],
            ["证据阻断", f"{data['totals']['blockingProcessEvidence']} 条 L3、{data['totals']['blockingBehaviorEvidence']} 条 A1 标记为缺原文证据；其中编号—名称不唯一分别为 {data['totals']['ambiguousProcessTitles']} / {data['totals']['ambiguousBehaviorTitles']} 条"],
        ]
    ):
        summary.rows[row].cells[0].text = values[0]
        summary.rows[row].cells[1].text = values[1]
        set_cell_shading(summary.rows[row].cells[0], SAGE)
    set_table_widths(summary, [1.25, 5.25])
    style_table(summary, header=False)

    add_callout(
        document,
        "本周完成标准",
        "各部门完善流程和业务行为描述，并至少选取 3 条流程完成字段级数据字典。所有流程、A1 和字段都要能回到原文制度或实际资料。",
        LIGHT_YELLOW,
    )
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("编制日期：2026-07-17\n").bold = True
    p.add_run(f"流程映射快照：{data['snapshotDate']}｜预填内容均需部门确认")
    document.add_page_break()

    document.add_heading("1. 本周要交付什么", level=1)
    add_bullet(document, "先在“01_流程总览”确认流程范围、系统接管期待和流程级边界。")
    add_bullet(document, "再在“02_业务行为”把每个 A1 写成可执行、可判断、可验收的业务动作。")
    add_bullet(document, "每部门至少选 3 条流程，在“03_数据字典”中按一字段一行拆解。")
    add_bullet(document, "所有流程、行为和字段都必须显示原文资料名称；缺证时保留“缺原文证据”，不得补写未经原文支持的结论。")
    add_callout(document, "责任边界", "信息化负责提供结构、检查规则和汇总方法；流程责任、制度解释和字段业务含义由所属部门确认。")

    document.add_heading("2. 八项要求：怎样把流程说清楚", level=1)
    questions = [
        ["1", "解决什么事", "流程目的和结束边界", "01_流程总览"],
        ["2", "谁负总责", "整条流程的责任角色，不是列出所有参与者", "01_流程总览"],
        ["3", "何时触发", "业务场景、事件或周期", "01/02"],
        ["4", "开始前有什么", "前置条件和输入材料", "01/02"],
        ["5", "每一步谁来做", "保留制度原文中的角色称谓", "02_业务行为"],
        ["6", "具体做什么", "动作、判断条件、下一步和退回位置", "02_业务行为"],
        ["7", "按什么标准", "时限、执行规则和验收条件", "02_业务行为"],
        ["8", "最终交付什么", "输出结果和完成标志", "01/02"],
    ]
    add_table(document, ["序号", "必须回答", "填写重点", "主要位置"], questions, [0.48, 1.2, 3.65, 1.17])

    document.add_heading("3. 完整示例：从 L3 到 A1", level=1)
    process, behaviors = find_example(data)
    add_callout(document, "示例原文出处", process["citationDisplay"] + "。下列内容来自当前流程映射，均须由财务部确认；责任角色归纳单独标记为待确认。", LIGHT_BLUE)
    process_rows = [
        ["1 解决什么事", process["purposeAndBoundary"]],
        ["2 谁负总责", "财务部成本会计（由 A1 执行角色归纳，待部门确认）"],
        ["3 何时触发", process["overallTrigger"]],
        ["4 开始前有什么", process["startConditionsAndInputs"]],
        ["5 每一步谁来做", "财务部成本会计；各 A1 继续逐行保留原文角色称谓"],
        ["6 具体做什么", "计算工序成本 → 计算废品损失 → 计算零件制造成本 → 分摊期间费用并计算全成本"],
        ["7 按什么标准", "过程计算规则按各 A1 对应条款填写；最终验收以全成本测算结果已传递至经营发展部为准"],
        ["8 最终交付什么", process["finalDeliverableAndCompletion"]],
    ]
    add_table(document, ["八项要求", "示例填写"], process_rows, [1.35, 5.15])
    add_label(document, "A1 行为示例（原文制度名称必须在主表中直接可见）")
    behavior_rows = []
    for item in behaviors[:4]:
        behavior_rows.append([
            item["a1Name"],
            item["actor"],
            item["triggerScene"],
            item["precondition"],
            item["outputResult"],
            item["citationDisplay"],
        ])
    add_table(document, ["A1业务行为", "执行角色", "触发", "前置", "输出", "原文出处"], behavior_rows, [1.4, 0.85, 1.05, 1.2, 0.85, 1.15])

    document.add_heading("4. 正常路径、判断分支与退回路径", level=1)
    workflow_image = asset_dir / "产品制造大纲工作流.png"
    if workflow_image.exists():
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(workflow_image), width=Inches(6.45))
        caption = document.add_paragraph("图 1  勤哲“产品制造大纲工作流”截图（示例资料，非公司制度）")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].italic = True
        caption.runs[0].font.size = Pt(8.5)
    add_number(document, "先写正常路径：编制 → 校对 → 审核 → 质保 → 无损检测 → 批准。")
    add_number(document, "判断节点单独成行：是否需要无损检测；需要则进入“无损检测”，不需要则直接进入“批准”。")
    add_number(document, "每条不同意路径都写明退回位置：校对、审核、质保、无损检测、批准不同意均退回“编制”。")
    add_number(document, "不要只写“退回”或“审批不通过”；必须写清条件、去向和重新进入正常路径的位置。")
    add_callout(document, "制度追溯写法", "即使流程结构来自截图，正式填报时仍要写“制度编号＋《原文制度名称》＋原文位置”。截图只能作为现场资料证据，不能替代制度原文。")

    document.add_heading("5. 原文制度和证据怎样填", level=1)
    evidence_rows = [
        ["制度直接证明本行为", "本行为直接引用", "GLC120110《设计评审管理程序》§5.2"],
        ["A1只有编号和条款，编号可唯一匹配", "补出完整原文名称，同时保留原始编号和条款", "GLTX-CW-06-A《成本测算管理程序》§5.5"],
        ["A1来自上下文推断", "继承所属流程制度，非本行为直接证据", "仍显示流程制度名称，并在引用方式中明示继承"],
        ["一条对象有多份制度", "主表用分号并列；04 证据索引一份来源一行", "不得只保留第一份"],
        ["只有表单、台账、截图或现场资料", "填写实际资料名称并标记未提供制度原文", "不得虚构制度名称"],
        ["制度标题无法唯一匹配", "缺原文证据", "属于阻断问题，不能标为已完成"],
    ]
    add_table(document, ["场景", "处理方式", "展示要求"], evidence_rows, [1.75, 2.8, 1.95])
    add_callout(document, "禁止项", "禁止只写“GLC120110 §5.2”而不显示《设计评审管理程序》；禁止把文件名当作制度标题；禁止为缺证记录编造制度名称。", LIGHT_YELLOW)

    document.add_heading("6. 数据字典：一字段一行", level=1)
    add_bullet(document, "每个 A1 已预置一条“待拆字段”起始行。现有输入、输出只作为字段发现线索，不代表已经完成字段拆解。")
    add_bullet(document, "本周先选至少 3 条流程；覆盖业务输入、系统自动字段、隐藏字段、枚举字段和审批留痕字段。")
    add_bullet(document, "字段必须关联到具体 L3 和 A1，并同时填写字段证据来源的原文资料名称、编号和位置。")
    field_rows = [
        ["制造大纲编号", "manufacturing_outline_no", "文本", "100", "是", "是", "是", "否", "勤哲产品制造大纲_主表字段截图（示例资料，非公司制度）"],
        ["表编号", "form_no", "文本", "100", "是", "否", "是", "是", "勤哲产品制造大纲_主表字段截图（示例资料，非公司制度）"],
        ["编制", "prepared_by", "人员", "20", "是", "否", "是", "是", "勤哲产品制造大纲_主表字段截图（示例资料，非公司制度）"],
        ["材料名称", "material_name", "文本", "100", "是", "否", "是", "否", "勤哲材料页明细表截图（示例资料，非公司制度）"],
        ["工序号", "operation_no", "文本", "20", "待确认", "否", "是", "否", "勤哲工序页明细表截图（示例资料，非公司制度）"],
        ["工时", "work_hours", "小数", "—", "待确认", "否", "否", "待确认", "勤哲工序页明细表截图（示例资料，非公司制度）"],
    ]
    add_table(document, ["字段中文名", "候选英文名", "类型", "长度", "必填", "主键", "查询", "自动", "原文资料名称"], field_rows, [0.82, 1.02, 0.58, 0.45, 0.48, 0.48, 0.48, 0.48, 1.71])
    document.add_paragraph()
    add_callout(document, "候选英文名", "英文名只作为后续系统设计候选，不在本周替代部门对字段中文名、业务定义和数据来源的确认。")
    add_label(document, "字段属性补充示例")
    attribute_rows = [
        ["表编号", "主表", "可见 / 不可编辑 / 不隐藏", "自动编号", "否", "勤哲产品制造大纲_主表字段截图（示例资料，非公司制度）"],
        ["是否需要无损检测", "主表", "可见 / 可编辑 / 不隐藏", "枚举：是 / 否", "作为流程分支判断", "勤哲产品制造大纲_主表字段截图（示例资料，非公司制度）"],
        ["批准", "主表", "可见 / 不可编辑 / 不隐藏", "当前用户姓名", "是，记录批准人和日期", "勤哲产品制造大纲_主表字段截图（示例资料，非公司制度）"],
        ["材料名称", "明细表", "可见 / 可编辑 / 不隐藏", "否", "否", "勤哲材料页明细表截图（示例资料，非公司制度）"],
        ["工序号", "明细表", "可见 / 可编辑 / 不隐藏", "否", "否", "勤哲工序页明细表截图（示例资料，非公司制度）"],
        ["内部流程实例ID（候选）", "主表", "不可见 / 不可编辑 / 隐藏", "自动生成", "是", "未提供原文资料；仅为隐藏控制字段写法示例，缺原文证据"],
    ]
    add_table(document, ["字段", "主表/明细", "显示、编辑、隐藏", "自动/枚举", "审批留痕", "原文资料名称"], attribute_rows, [0.95, 0.7, 1.3, 0.85, 1.05, 1.65])

    document.add_heading("7. 八张工作表怎么用", level=1)
    sheets = [
        ["00_填写说明", "看数量、颜色、八项要求和填写顺序"],
        ["01_流程总览", "确认 L3、制度、系统承接方向、信息化接管期待和流程级边界"],
        ["02_业务行为", "逐条完善 A1 的角色、触发、输入、动作、判断、标准、输出和完成标志"],
        ["03_数据字典", "一字段一行；每部门本周至少完成 3 条流程"],
        ["04_证据索引", "保存一对多的制度、表单、台账、流程图和现场资料关系"],
        ["05_完整性检查", "查看缺制度名、缺位置、八项不完整、判断路径不完整和字段覆盖情况"],
        ["98_下拉选项", "标准状态和字段属性选项；不删除已有值"],
        ["99_来源快照", "生成时原始映射，只读参考，用于对比部门调整"],
    ]
    add_table(document, ["工作表", "用途"], sheets, [1.45, 5.05])
    add_callout(document, "颜色", "灰色为来源预填，浅黄色为部门填写，浅蓝色为公式结果。状态只使用“已完成 / 待部门确认 / 缺原文证据”。", LIGHT_BLUE)

    document.add_heading("8. 部门提交前的评审清单", level=1)
    checks = [
        "每条 L3 和 A1 在主表中都能直接看到原文资料名称；没有只写制度编号。",
        "制度编号、制度名称、原始文件名和原文位置能够相互核对。",
        "流程级五项内容与 A1 级行为描述共同覆盖八项要求。",
        "判断节点有判断条件、正常去向和退回位置。",
        "每个 A1 都检查了输入字段、自动字段、隐藏字段、枚举字段和审批留痕字段。",
        "本部门至少 3 条流程的数据字典达到“已完成”。",
        "系统承接待确认流程仍保留，并填写是否期望信息化接管及接管范围。",
        "无法回到原文的内容保留为“缺原文证据”，未用经验替代制度。",
    ]
    for check in checks:
        add_bullet(document, "□ " + check)

    document.add_heading("9. 下周任务预告：把字段连成数据关系", level=1)
    next_week = [
        ["数据从哪里来", "来源部门、来源角色、来源表单/台账/系统、产生时点"],
        ["个人如何处理", "查看、录入、选择、计算、判断、审批、转交、归档等具体动作"],
        ["处理产生什么结果", "字段值变化、中间表、审批意见、编号、状态或异常记录"],
        ["最终交付到什么状态", "交付对象、接收部门、完成标志、后续流程入口"],
    ]
    add_table(document, ["下周要回答", "梳理口径"], next_week, [1.65, 4.85])
    add_callout(document, "本周与下周的衔接", "本周先把流程、行为和字段说清楚；下周基于这些字段梳理来源、处理方式、处理结果和最终交付状态，形成可追溯的数据关系。", SAGE)

    add_header_footer(document, data["snapshotDate"])
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> None:
    args = parse_args()
    data = json.loads(Path(args.data).read_text(encoding="utf-8"))
    build_document(data, Path(args.output), Path(args.asset_dir))
    print(json.dumps({"output": args.output, "snapshotDate": data["snapshotDate"], "processes": data["totals"]["processes"], "behaviors": data["totals"]["behaviors"]}, ensure_ascii=False))


if __name__ == "__main__":
    main()
